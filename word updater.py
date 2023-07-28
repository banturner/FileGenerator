import docx

def update_placeholders(document, placeholders):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in placeholders.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in placeholders.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

def main():
    template_filename = "template.docx"
    output_filename = "output.docx"

    name = input("Enter name: ")
    date = input("Enter date(DD/MM/YY): ")
    serial_number = input("Enter serial number: ")
    company_name = input("Enter company name: ")
    id = input("Enter customer ID: ")
    Tin = input("Enter time in (24Hr): ")
    Tout = input("Enter time out (24Hr): ")
    Temi = input("Enter temi version: ")
    DOW = input("Enter description of work: ")
    pn = input("Select the part number: ")
    pd = input("Select the part description: ")
    qty = input("Enter the quantity: ")
    price = input("Select NA or As Quoted: ")
    comments = input("Enter the comments: ")


    placeholders = {
        "name": name,
        "date": date,
        "serial": serial_number,
        "company": company_name,
        "id": id,
        "tin": Tin,
        "tout": Tout,
        "Temi": Temi,
        "DOW": DOW,
        "pn":pn,
        "pd":pd,
        "qty":qty,
        "price":price,
        "comments":comments,

    }

    document = docx.Document(template_filename)
    update_placeholders(document, placeholders)
    document.save(output_filename)

if __name__ == "__main__":
    main()
