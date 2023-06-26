from docx import Document

def generate_word_document(filename, data):
    # Create a new Word document
    document = Document()

    # Write data to the document
    for row_data in data:
        row = document.add_paragraph()
        for cell_data in row_data:
            row.add_run(cell_data)

    # Save the document
    document.save(filename)
    print(f"Word document '{filename}' generated successfully.")

# Example usage
data = [
    ["Heading 1", "Heading 2", "Heading 3"],
    ["Value 1", "Value 2", "Value 3"],
    ["Value 4", "Value 5", "Value 6"]
]

filename = "example.docx"
generate_word_document(filename, data)
